"""Create a sample .docx file for demo purposes."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import os

skill_dir = Path(os.environ.get("SKILL_DIR", Path(__file__).parent))

doc = Document()

# Title
title = doc.add_heading("Skills Sandbox Demo", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# Intro
doc.add_paragraph(
    "This document was created inside a Docker sandbox by the docx-to-pdf skill. "
    "It demonstrates that the sandbox environment has all required dependencies "
    "pre-installed and working."
)

# What this proves
doc.add_heading("What this proves", level=2)
items = [
    "python-docx is installed and can create Word documents",
    "LibreOffice is installed and can convert DOCX → PDF",
    "The sandbox runs as a non-root user (sandbox, UID 1000)",
    "Skills run identically for every teammate via Docker",
    "The environment is reproducible from a single Dockerfile",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

# How it works
doc.add_heading("How it works", level=2)
steps = [
    ("Git repo", "Stores the Dockerfile, requirements.txt, and all skills"),
    ("Docker image", "Contains Python, Node.js, LibreOffice, and all dependencies"),
    ("make run", "Launches a container, mounts the skills, runs the entry point"),
    ("GitHub Actions", "Auto-builds the image and runs tests on every push"),
]
table = doc.add_table(rows=1, cols=2, style="Light Shading Accent 1")
table.rows[0].cells[0].text = "Component"
table.rows[0].cells[1].text = "Purpose"
for component, purpose in steps:
    row = table.add_row()
    row.cells[0].text = component
    row.cells[1].text = purpose

doc.add_paragraph("")
doc.add_paragraph(
    "Generated automatically — no manual installation required.",
)

output = skill_dir / "demo_input.docx"
doc.save(str(output))
print(f"[docx-to-pdf] Created: {output.name}")
